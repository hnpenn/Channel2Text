# PBPScrapeBot.py
import os

import discord

#import nest_asyncio
#nest_asyncio.apply()

from dotenv import load_dotenv

from discord.ext import commands

import csv
import docx

import demoji

load_dotenv()
TOKEN = os.getenv('DISCORD_TOKEN')

bot = commands.Bot(command_prefix='+')

def get_data(messageRead):
    mContent = messageRead.content
    mContent = demoji.replace(mContent, repl = "")
    mAuthor = str(messageRead.author.name)
    mDate = messageRead.created_at.strftime("%Y-%m-%d, %H:%M:%S")
    return [mDate, mAuthor, mContent]

async def csv_all(ctx, channelUse):
    counter = 0
    async with ctx.channel.typing():
        with open('test.csv', mode='w', newline='') as message_file:
            message_writer = csv.writer(message_file, delimiter=',', quotechar='"')
            message_writer.writerow(['Date', 'Author', 'Message'])
            async for message in channelUse.history(limit = None):
                message_writer.writerow(get_data(message))
                counter += 1
            await ctx.send("Message total: " + str(counter))
            
async def csv_me(ctx, channelUse, user_me):
    counter = 0
    async with ctx.channel.typing():
        with open('test.csv', mode='w', newline='') as message_file:
            message_writer = csv.writer(message_file, delimiter=',', quotechar='"')
            message_writer.writerow(['Date', 'Author', 'Message'])
            async for message in channelUse.history(limit = None):
                if message.author == user_me:
                    message_writer.writerow(get_data(message))
                    counter += 1
            await ctx.send("Message total: " + str(counter))
            
def index_list(sentences, symbol):
    #print(symbol)
    idxSet = set()
    idx = 0
    pr_idx = -1
    while idx != -1:
        idx = sentences.find(symbol, pr_idx + 1, len(sentences))
        #print("idx: " + str(idx))
        #print(sentences)
        idxSet.add(idx)
        pr_idx = idx
    idxSet.remove(-1)
    temp_list = list(idxSet)
    if (len(temp_list)) % 2 != 0:
        temp_list.sort()
        temp_list.pop() 
    idxSet = set(temp_list)
    return idxSet

def to_format_dict(idxSet, symbol):
    temp_list = list(idxSet)
    idx_dict = {}
    for idx in temp_list:
        idx_dict.update({idx: symbol})
    return idx_dict
            
def docx_format(para, sentences):
    us_set = index_list(sentences, '_')
    doub_set = index_list(sentences, '**')
    sing_set = index_list(sentences, '*')
    trip_set = index_list(sentences, '***')
    
    print('_')
    print(us_set)
    print('**')
    print(doub_set)
    print('*')
    print(sing_set)
    print('***')
    print(trip_set)
    
    doub_set.difference_update(trip_set)
    sing_set.difference_update(doub_set)
    doub_list_2 = list(doub_set)
    doub_set_2 = set()
    for idx in doub_list_2:
        idx2 = idx + 1
        doub_set_2.add(idx2)
    sing_set.difference_update(doub_set_2)
    
    if (len(us_set) + len(doub_set) + len(sing_set)) == 0:
        para = para.insert_paragraph_before(sentences)
        return para
    
    us_dict = to_format_dict(us_set, '_')
    doub_dict = to_format_dict(doub_set, '**')
    sing_dict = to_format_dict(sing_set, '*')
    
    format_dict = {}
    format_dict.update(us_dict)
    format_dict.update(doub_dict)
    format_dict.update(sing_dict)
    
    idx_keys = format_dict.keys()
    idx_list = list(idx_keys)
    idx_list.sort()
    
    format_key = {'_': False, '*': False, '**': False}
    
    paragraph = para.insert_paragraph_before()
    
    pr_idx = -1
    for idx in idx_list:
        txt = sentences[(pr_idx + 1):idx]
        txt = txt.strip('*_')
        run = paragraph.add_run(txt)
        if format_key['**']:
            run.bold = True
        if format_key['*'] or format_key['_']:
            run.italic = True
        symbol = format_dict[idx]
        if format_key[symbol]:
            format_key.update({symbol: False})
        else:
            format_key.update({symbol: True})
        pr_idx = idx
    if pr_idx != (len(sentences) - 1):
        paragraph.add_run(sentences[(pr_idx + 1):].strip('*_'))
    return paragraph
            
async def docx_all(ctx, channelUse):
    counter = 0
    async with ctx.channel.typing():
        doc = docx.Document()
        doc.save("test.docx")
        para = doc.add_paragraph("")
        async for message in channelUse.history(limit = None):
            sentences = message.clean_content
            if '*' in sentences or '_' in sentences:
                #print("message " + str(counter) + " is fancy")
                para = docx_format(para, sentences)
            else:
                #print("message " + str(counter) + " is plain")
                para = para.insert_paragraph_before(message.clean_content)
            counter += 1
        #print("I finished formatting")
        doc.save("test.docx")
        await ctx.send("Message total: " + str(counter))
            
async def docx_me(ctx, channelUse, user_me):
    counter = 0
    async with ctx.channel.typing():
        doc = docx.Document()
        doc.save("test.docx")
        para = doc.add_paragraph("")
        async for message in channelUse.history(limit = None):
            if message.author == user_me:
                sentences = message.clean_content
                if '*' in sentences or '_' in sentences:
                #print("message " + str(counter) + " is fancy")
                    para = docx_format(para, sentences)
                else:
                    #print("message " + str(counter) + " is plain")
                    para = para.insert_paragraph_before(message.clean_content)
                counter += 1
        doc.save("test.docx")
        await ctx.send("Message total: " + str(counter))
    
async def txt_all(ctx, channelUse):
    counter = 0
    async with ctx.channel.typing():
        with open('test.txt', mode='w', newline='') as message_file:
            async for message in channelUse.history(limit = None):
                message_file.write(message.content)
                message_file.write("\n\n")
                counter += 1
            await ctx.send("Message total: " + str(counter))
            
async def txt_me(ctx, channelUse, user_me):
    counter = 0
    async with ctx.channel.typing():
        with open('test.txt', mode='w', newline='') as message_file:
            async for message in channelUse.history(limit = None):
                if message.author == user_me:
                    message_file.write(message.content)
                    message_file.write("\n\n")
                    counter += 1
            await ctx.send("Message total: " + str(counter))

async def data_all(ctx, channelUse, channelID, docType):
    response = "Getting all messages"
    await ctx.send(response)
    if docType == 'csv':
        await csv_all(ctx, channelUse)
    elif docType == 'docx':
        await docx_all(ctx, channelUse)
    else:
        await txt_all(ctx, channelUse)
    f = open('test.' + docType, 'rb')
    sendFile = discord.File(f, filename=channelID + "_transcript." + docType, spoiler = False)
    f.close()
    await ctx.send(content = "Special delivery", file = sendFile)
    
async def data_me(ctx, channelUse, channelID, docType):
    user_me = ctx.message.author
    response = "Getting messages for " + str(user_me)
    await ctx.send(response)
    if docType == 'csv':
        await csv_me(ctx, channelUse, user_me)
    elif docType == 'docx':
        await docx_me(ctx, channelUse, user_me)
    else:
        await txt_me(ctx, channelUse, user_me)
    f = open('test.' + docType, 'rb')
    sendFile = discord.File(f, filename=channelID + "_" + str(user_me) + "_transcript." + docType, spoiler = False)
    await ctx.send(content = "Special delivery", file = sendFile)
    f.close()

@bot.command(name = 'data')
async def data_collect(ctx, channelID, user, docType):
    docTypeList = ['csv', 'docx', 'txt']
    guildScrape = ctx.guild
    guildChannels = guildScrape.channels
    inGuild = False
    channelUse = None
    for channel in guildChannels:
        if channel.name == channelID:
            inGuild = True
            channelUse = channel
            break
    if not inGuild:
        response = "Error: Channel not in server. Remember the dashes!"
        await ctx.send(response)
        return
    if docType not in docTypeList:
        response = "Error: Please choose either a csv, docx, or txt file format.\n"
        await ctx.send(response)
        return
    if user == 'all':
        await data_all(ctx, channelUse, channelID, docType)
        return
    elif user == 'me':
        await data_me(ctx, channelUse, channelID, docType)
        return
    else:
        response = "Please, good sir, ma'am, or gentleperson. Use \"all\" or \"me\"."
        await ctx.send(response)
        return
        
    

@bot.event
async def on_ready():
    print(f'{bot.user.name} has connected to Discord!')

bot.run(TOKEN)
